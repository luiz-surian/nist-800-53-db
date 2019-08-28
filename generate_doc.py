#!/usr/bin/python
# -*- coding: utf-8 -*-

"""NIST 800-53 Controls - Word Document Generator"""

__title__ = 'NIST 800-53'
__author__ = 'Luiz Fernando Surian Filho'
__version__ = '1'
__language__ = 'en'
__license__ = 'GPL'

import os
import sqlite3
import uuid
from datetime import datetime

from docx import Document
from docx.shared import Pt


class MissingLanguage(LookupError):
    """Raise this when there's a missing language."""


def build_doc(target):
    if target == 'nist_800_53_en':
        language = {
            'control_description': 'Control Description',
            'supplemental_guidance': 'Supplemental Guidance',
            'control_enhancements': 'Control Enhancements',
            'references': 'References',
            'metadata': {
                'category': 'Security Controls List',
                'comments': 'Security Controls and Assessment Procedures for '
                            'Federal Information Systems and Organizations',
                'description': 'National Vulnerability Database',
                'language': 'en',
                'subject': 'This database represents the security controls and associated assessment procedures '
                           'defined in NIST SP 800-53 Revision 4 - Recommended Security Controls for '
                           'Federal Information Systems and Organizations.'
            }
        }
    elif target == 'nist_800_53_br':
        language = {
            'control_description': 'Descrição do Controle',
            'supplemental_guidance': 'Guia Adicional',
            'control_enhancements': 'Melhorias do Controle',
            'references': 'Referências',
            'metadata': {
                'category': 'Lista de Controles de Segurança',
                'comments': 'Controles de Segurança e Procedimentos de Avaliação para '
                            'Sistemas de Informação Federais e Organizaçionais',
                'description': 'Base de Dados de Vulnerabilidades',
                'language': 'pt-BR',
                'subject': 'Essa base de dados representa os controles de segurança e procedimentos de avaliação '
                           'definidos no NIST SP 800-53 Revisão 4 - Controles de Segurança Recomendados para '
                           'Sistemas de Informação Federais e Organizacionais.'
            }
        }
    else:
        raise MissingLanguage('This language was not implemented yet.')

    connection = sqlite3.connect(':memory:')
    cursor = connection.cursor()

    with open(f'database/{target}.sql', 'rb') as db_file:
        sequel = db_file.read().decode('utf-8')
    cursor.executescript(sequel)
    connection.commit()

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    document.add_heading('NIST 800-53', level=0)

    cursor.execute('SELECT * FROM control_families;')
    families = cursor.fetchall()

    for family_id, family_abbreviation, family_name in families:
        document.add_heading(f'{family_abbreviation} - {family_name}\n', level=1)

        cursor.execute(f'SELECT * FROM security_control WHERE family = {family_id};')
        controls = cursor.fetchall()
        for control_id, control_family, control_number, control_priority, control_impact, \
                control_name, control_description, control_guidance in controls:
            control_title = f'{family_abbreviation}-{control_number}: {control_name} (P{control_priority})'
            document.add_heading(control_title, level=2)
            paragraph = document.add_paragraph()
            paragraph.add_run(f'{language["control_description"]}:\n').bold = True
            paragraph.add_run(f'{control_description}\n')
            paragraph.add_run('\n')
            if control_guidance is not None:
                paragraph.add_run(f'{language["supplemental_guidance"]}:\n').bold = True
                paragraph.add_run(f'{control_guidance}\n')
                paragraph.add_run('\n')

            cursor.execute(f'SELECT * FROM control_enhancements WHERE control_id = {control_id};')
            enhancements = cursor.fetchall()
            if enhancements:
                paragraph.add_run(f'{language["control_enhancements"]}:\n').bold = True
                for enhancement_id, enhancement_control, enhancement_number, enhancement_impact, \
                        enhancement_name, enhancement_text, enhancement_supplemental in enhancements:
                    enhancement_title = f'{family_abbreviation}-{control_number}({enhancement_number}): ' \
                                        f'{control_name} | {enhancement_name}'
                    document.add_heading(enhancement_title, level=3)
                    paragraph = document.add_paragraph()
                    paragraph.add_run(f'{enhancement_text}\n')
                    if enhancement_supplemental:
                        paragraph.add_run(f'{language["supplemental_guidance"]}: {enhancement_supplemental}\n')

            cursor.execute(f'SELECT * FROM control_references WHERE control_id = {control_id};')
            references = cursor.fetchall()
            if references:
                paragraph = document.add_paragraph()
                paragraph.add_run(f'{language["references"]}:\n').bold = True
                table = document.add_table(rows=len(references), cols=2)
                i = 0
                for _, _, reference_id in references:
                    cursor.execute(f'SELECT * FROM references_links WHERE id = {reference_id};')
                    reference_link = cursor.fetchone()
                    table.rows[i].cells[0].text = reference_link[1]
                    table.rows[i].cells[1].text = reference_link[2]
                    i += 1

        document.add_page_break()

    today = datetime.now()
    meta = language['metadata']
    core_properties = document.core_properties
    core_properties.author = __author__
    core_properties.category = meta['category']
    core_properties.comments = meta['comments']
    core_properties.content_status = 'Final'
    core_properties.created = today
    core_properties.description = meta['description']
    core_properties.identifier = str(uuid.uuid4())
    core_properties.keywords = 'Banco Bradesco S.A.'
    core_properties.language = meta['language']
    core_properties.last_modified_by = __author__
    core_properties.last_printed = today
    core_properties.modified = today
    core_properties.revision = 1
    core_properties.subject = meta['subject']
    core_properties.title = __title__
    core_properties.version = 'Rev. 4'

    document.save(f'documents/{target}.docx')
    connection.close()


if __name__ == '__main__':
    for file in os.listdir('database'):
        filename = os.path.splitext(file)[0]
        build_doc(target=filename)
